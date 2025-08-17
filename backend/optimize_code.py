from fastapi import FastAPI, Form
from fastapi.responses import JSONResponse
from docx import Document
from docx2pdf import convert
from email.message import EmailMessage
import shutil
import os
import smtplib
import uvicorn
import uuid
from datetime import datetime
import yaml
import re

app = FastAPI()

# ---------------- CONFIG ----------------
TEMPLATE_PATH = "C:/Users/Akash Verma/Desktop/akash_learning/python/projects/registration_recipt_generator/registration_successful_recept.docx"
OUTPUT_DIR = "students_registration_recipt"
MASTER_DATA_PATH = "master_data"
MASTER_DATA_FILE = os.path.join(MASTER_DATA_PATH, "registrations.yaml")

os.makedirs(MASTER_DATA_PATH, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Use environment variables for email credentials
SENDER_EMAIL = os.getenv("SENDER_EMAIL", "your_email@gmail.com")
SENDER_PASSWORD = os.getenv("SENDER_PASSWORD", "rkvb nfwk ixqz yxjy")


# ---------------- UTILITIES ----------------
def load_master_data():
    if os.path.exists(MASTER_DATA_FILE):
        with open(MASTER_DATA_FILE, "r") as f:
            return yaml.load(f, Loader=yaml.FullLoader) or []
    return []


def save_master_data(data):
    with open(MASTER_DATA_FILE, "w") as f:
        yaml.dump(data, f)


def generate_application_no(timestamp):
    data = load_master_data()
    return f"{timestamp}{len(data)+1}"


def is_transaction_exists(transaction_id):
    data = load_master_data()
    return any(entry["transaction_id"] == transaction_id for entry in data)


def sanitize_filename(name):
    """Remove spaces and special characters for safe file names"""
    return re.sub(r"[^A-Za-z0-9_-]", "", name.replace(" ", "_"))


def generate_docx(template_path, output_path, replacements: dict):
    shutil.copy(template_path, output_path)
    doc = Document(output_path)

    # Replace in paragraphs
    for p in doc.paragraphs:
        for key, val in replacements.items():
            if key in p.text:
                p.text = p.text.replace(key, str(val))

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(val))

    doc.save(output_path)


def convert_to_pdf(input_docx, output_pdf):
    convert(input_docx, output_pdf)


def send_email(recipient_email, subject, body, attachment_path):
    try:
        msg = EmailMessage()
        msg["From"] = SENDER_EMAIL
        msg["To"] = recipient_email
        msg["Subject"] = subject
        msg.set_content(body)

        with open(attachment_path, "rb") as f:
            file_data = f.read()
            file_name = os.path.basename(attachment_path)
        msg.add_attachment(
            file_data, maintype="application", subtype="pdf", filename=file_name
        )

        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(SENDER_EMAIL, SENDER_PASSWORD)
            server.send_message(msg)
    except Exception as e:
        raise RuntimeError(f"Failed to send email: {e}")


# ---------------- API ENDPOINTS ----------------
@app.get("/healthz/")
async def health_check():
    return {"status": "ok"}


@app.post("/generate-receipt/")
async def generate_receipt(
    student_name: str = Form(...),
    branch: str = Form(...),
    year: str = Form(...),
    college: str = Form(...),
    mobile: str = Form(...),
    email: str = Form(...),
    course: str = Form(...),
    pay_for: str = Form(...),
    amount: str = Form(...),
    payment_mode: str = Form(...),
    transaction_id: str = Form(...),
    payment_date: str = Form(...),
):
    # Check if transaction exists
    if is_transaction_exists(transaction_id):
        return JSONResponse(
            status_code=400,
            content={"error": f"Transaction ID {transaction_id} already exists"},
        )

    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    receipt_no = f"CTC{timestamp}{str(uuid.uuid4())[:3]}"
    application_no = generate_application_no(timestamp)

    # Create entry
    entry = {
        "receipt_no": receipt_no,
        "application_no": application_no,
        "student_name": student_name.upper(),
        "branch": branch,
        "year": year,
        "college": college,
        "course": course,
        "mobile": mobile,
        "email": email,
        "pay_for": pay_for,
        "amount": amount,
        "payment_mode": payment_mode,
        "payment_date": payment_date,
        "transaction_id": transaction_id,
        "timestamp": timestamp,
    }

    # Save entry
    data = load_master_data()
    data.append(entry)
    save_master_data(data)

    # File paths
    safe_name = sanitize_filename(student_name)
    filled_docx = os.path.join(OUTPUT_DIR, f"{safe_name}_{receipt_no}.docx")
    output_pdf = os.path.join(OUTPUT_DIR, f"{safe_name}_{receipt_no}.pdf")

    # Generate documents
    replacements = {
        "{{receipt_no}}": receipt_no,
        "{{application_no}}": application_no,
        "{{department}}": branch,
        "{{student_name}}": student_name.upper(),
        "{{course_name}}": course,
        "{{payment_mode}}": payment_mode,
        "{{pay_for}}": pay_for,
        "{{payment_date}}": payment_date,
        "{{amount}}": amount,
        "{{transaction_id}}": transaction_id,
    }

    try:
        generate_docx(TEMPLATE_PATH, filled_docx, replacements)
        convert_to_pdf(filled_docx, output_pdf)
        send_email(
            email,
            "Your Fee Payment Receipt",
            f"Dear {student_name},\n\nPlease find attached your payment receipt.\n\nThanks.",
            output_pdf,
        )
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})

    return JSONResponse(
        {
            "message": "Receipt generated, saved, and sent successfully!",
            "receipt_no": receipt_no,
            "application_no": application_no,
            "docx_path": filled_docx,
            "pdf_path": output_pdf,
        }
    )


if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
