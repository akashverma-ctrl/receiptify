from fastapi import FastAPI, Form
from fastapi.responses import JSONResponse
from docx import Document
from docx2pdf import convert
from fastapi.middleware.cors import CORSMiddleware

import shutil
import os
import smtplib
from email.message import EmailMessage
import uvicorn
import uuid
from datetime import datetime
import yaml

# import Exception

app = FastAPI()

TEMPLATE_PATH = "C:/Users/Akash Verma/Desktop/akash_learning/python/projects/registration_recipt_generator/backend/registration_successful_recept.docx"
OUTPUT_DIR = "students_registration_recipt"
MASTER_DATA_PATH = "master_data"
MASTER_DATA_FILE = os.path.join(MASTER_DATA_PATH, "registrations.yaml")
os.makedirs(MASTER_DATA_PATH, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "*"
    ],  # or ["http://127.0.0.1:5500"] if serving frontend with Live Server
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def update_master_file(data: list):
    with open(MASTER_DATA_FILE, "w") as f:
        yaml.dump(data, f)


def generate_application_number(entry: dict):
    """Append a row in registrations.xlsx, create file if not exists"""
    data = []
    if os.path.exists(MASTER_DATA_FILE):
        with open(MASTER_DATA_FILE, "r") as f:
            data = yaml.load(f, Loader=yaml.FullLoader) or []
    application_number = entry.get("timestamp") + str(len(data) + 1)
    entry["application_no"] = application_number
    data = data + [entry]

    return application_number, data


def is_user_and_payment_exists(transaction_id):
    data = []
    if os.path.exists(MASTER_DATA_FILE):
        with open(MASTER_DATA_FILE, "r") as f:
            data = yaml.load(f, Loader=yaml.FullLoader) or []
    for entry in data:
        if entry.get("transaction_id") == transaction_id:
            return True
    return False


def generate_docx_from_template(template_path, output_path, replacements: dict):
    shutil.copy(template_path, output_path)
    doc = Document(output_path)

    for p in doc.paragraphs:
        for key, val in replacements.items():
            if key in p.text:
                p.text = p.text.replace(key, str(val))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(val))

    doc.save(output_path)


def convert_to_pdf(input_docx, output_pdf):
    convert(input_docx, output_pdf)


def send_email(
    sender_email, sender_password, recipient_email, subject, body, attachment_path
):
    msg = EmailMessage()
    msg["From"] = sender_email
    msg["To"] = recipient_email
    msg["Subject"] = subject
    msg.set_content(body)

    with open(attachment_path, "rb") as f:
        file_data = f.read()
        file_name = os.path.basename(attachment_path)
    msg.add_attachment(
        file_data, maintype="application", subtype="pdf", filename=file_name
    )

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
        server.login(sender_email, sender_password)
        server.send_message(msg)


@app.get("/healthz/")
async def health_check():
    return {"status": "ok"}


@app.post("/generate-receipt/")
async def generate_receipt(
    student_name: str = Form(...),
    branch: str = Form(...),
    year: str = Form(...),
    college: str = Form(...),
    mobile: str = Form(...),
    email: str = Form(...),
    course: str = Form(...),
    pay_for: str = Form(...),
    amount: str = Form(...),
    payment_mode: str = Form(...),
    transaction_id: str = Form(...),
    payment_date: str = Form(...),
):
    if is_user_and_payment_exists(transaction_id):
        return JSONResponse(
            {
                "status_code": 400,
                "error": True,
                "message": f"Transaction id {transaction_id} already exists",
            }
        )
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")

    # Prepare preliminary receipt_no and application_no
    receipt_no = "CTC" + timestamp + str(uuid.uuid4())[:3]

    # Create an entry dict (application_no filled later after Excel row is known)
    entry = {
        "receipt_no": receipt_no,
        "student_name": student_name.upper(),
        "branch": branch,
        "year": year,
        "college": college,
        "course": course,
        "mobile": mobile,
        "email": email,
        "pay_for": pay_for,
        "amount": amount,
        "payment_mode": payment_mode,
        "payment_date": payment_date,
        "transaction_id": transaction_id,
        "timestamp": timestamp,
    }

    # Save to yaml and get application_no
    application_no, master_data = generate_application_number(entry)
    entry["application_no"] = application_no

    # File paths
    safe_name = student_name.replace(" ", "_")
    filled_docx = os.path.join(OUTPUT_DIR, f"{safe_name}_{receipt_no}.docx")
    output_pdf = os.path.join(OUTPUT_DIR, f"{safe_name}_{receipt_no}.pdf")

    # Replacements for Word template
    replacements = {
        "{{receipt_no}}": entry["receipt_no"],
        "{{application_no}}": entry["application_no"],
        "{{department}}": entry["branch"],
        "{{student_name}}": entry["student_name"],
        "{{course_name}}": entry["course"],
        "{{payment_mode}}": entry["payment_mode"],
        "{{pay_for}}": entry["pay_for"],
        "{{payment_date}}": entry["payment_date"],
        "{{amount}}": entry["amount"],
        "{{transaction_id}}": entry["transaction_id"],
    }

    # Generate documents
    generate_docx_from_template(TEMPLATE_PATH, filled_docx, replacements)
    convert_to_pdf(filled_docx, output_pdf)

    # Send email
    send_email(
        sender_email="campustocorporate.academy@gmail.com",
        sender_password="rkvb nfwk ixqz yxjy",
        recipient_email=email,
        subject="Your Fee Payment Receipt",
        body=f"Dear {student_name},\n\nPlease find attached your payment receipt.\n\nThanks.",
        attachment_path=output_pdf,
    )

    # Save master data
    update_master_file(master_data)

    return JSONResponse(
        {
            "status_code": 200,
            "error": False,
            "message": f"Receipt generated, saved, and sent successfully! Your Receipt No. is : {receipt_no} and your Application No. is : {application_no}",
            "receipt_no": receipt_no,
            "application_no": application_no,
        }
    )


if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
